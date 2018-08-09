#coding=utf8
import os
import sys
import datetime as dt
import numpy as np
import pandas as pd

sys.path.append(r'D:\Anaconda3\Lib\site-packages')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import matplotlib as mpl
import matplotlib.pyplot as plt

from WindPy import w
import cx_Oracle
os.environ['NLS_LANG'] = 'SIMPLIFIED CHINESE_CHINA.UTF8'


def get_all_from_list(filePath=None):
    if filePath is None:
        filePath = r'.\holdLists'
    allLists = os.listdir(filePath)
    allDates = sorted([int(fl.split('.')[0].split('_')[-1]) for fl in allLists])
    cols = ['stkcd','buyDate','ret','flag']
    outPut = {
        'tradeDate': [],
        'buyReturn': [],
        'holdReturn': [],
        'netReturn':[]
    }
    for tdt in allDates:
        retList = pd.read_csv(os.path.join(filePath,'returns_tradeDate_{}.csv'.format(tdt)),encoding='gbk')
        retList = retList.loc[:, cols]
        validBuy = np.all([retList['buyDate']==tdt, np.isin(retList['flag'],(0,3,4))] ,axis=0)
        validHold = np.all([retList['buyDate']<tdt, np.isin(retList['flag'],(0,2,3,4))] ,axis=0)
        outPut['tradeDate'].append(tdt)
        outPut['buyReturn'].append(retList.loc[validBuy,'ret'].mean())
        outPut['holdReturn'].append(retList.loc[validHold,'ret'].mean())
        outPut['netReturn'].append(retList.loc[validBuy | validHold,'ret'].mean())
    outPut = pd.DataFrame(outPut).loc[:,['tradeDate','buyReturn','holdReturn','netReturn']]


def generate_return_report(tradeDate,tradeListPath=None,holdListPath=None):
    tradeDate = str(tradeDate)
    holdListPath = r'.\holdLists' if holdListPath is None else holdListPath
    tradeListPath = r'C:\Users\Administrator\Desktop\simu_report\buyLists' if tradeListPath is None else tradeListPath
    w.start()
    infoDate = w.tdaysoffset(-1,tradeDate).Data[0][0].strftime('%Y%m%d')
    # infoDate = '20180808'
    retList = pd.read_csv(os.path.join(holdListPath,'returns_tradeDate_{}.csv'.format(infoDate)),encoding='gbk')
    trdList = pd.read_csv(os.path.join(tradeListPath,'stkNum50_tradeList_infoDate_{}.csv'.format(infoDate)),encoding='gbk')
    holdStkList = retList.loc[retList['buyDate']==int(infoDate),['stkcd','predictVal']].set_index(['stkcd'])
    buyStkList = trdList.loc[:,['stkcd','predictVal']].set_index(['stkcd'])

    conn = cx_Oracle.connect(r'c##sensegain/sensegain@220.194.41.75/wind')
    cursor = conn.cursor()
    sqlLines = 'SELECT TRADE_DT, S_INFO_WINDCODE, S_DQ_OPEN, S_DQ_CLOSE, S_DQ_PRECLOSE, S_DQ_PCTCHANGE/100, S_DQ_CLOSE/S_DQ_OPEN-1, S_DQ_TRADESTATUS ' \
               'FROM c##wind.AShareEODPrices WHERE TRADE_DT={0}'.format(tradeDate)
    rets = cursor.execute(sqlLines).fetchall()
    rets = pd.DataFrame(rets, columns=['tradeDate', 'stkcd', 'open', 'close', 'preclose', 'retCC','retOC','trdStat'])
    rets['flagOC'] = 1*(rets['trdStat']=='停牌') + 2*(rets['open']/rets['preclose']>=1+0.099) + 3*(rets['open']/rets['preclose']<=1-0.099)
    rets['flagCC'] = 1*(rets['trdStat']=='停牌') + 2*(rets['retCC'] >= 0.099) + 3 * (rets['retCC'] <= -0.099)
    rets['stkcd'] = rets['stkcd'].map(lambda x:int(x[:6]))
    rets.set_index(['stkcd'],inplace=True)
    holdStkList = holdStkList.join(rets.loc[:,['tradeDate','close','retCC','flagCC']])
    holdStkList.columns = ['predictVal','date','close','ret','flag']
    holdStkList['buyDate'] = infoDate
    buyStkList = buyStkList.join(rets.loc[:, ['tradeDate', 'close', 'retOC', 'flagOC']])
    buyStkList.columns = ['predictVal','date','close','ret','flag']
    buyStkList['buyDate'] = tradeDate
    output = pd.concat([holdStkList,buyStkList],axis=0)
    output = output.loc[:,['date','close','ret','flag','buyDate','flag']]
    output.to_csv(os.path.join(holdListPath,'returns_tradeDate_{}.csv'.format(tradeDate)))
    print('return report generated for trade date {}'.format(tradeDate))


def update_recorder(endDate=None,recordFile=None,holdListPath=None):
    endDate = dt.datetime.today().strftime('%Y%m%d') if endDate is None else endDate
    recordFile = r'.\backtest_recorder.csv' if recordFile is None else recordFile
    holdListPath = r'.\holdLists' if holdListPath is None else holdListPath
    recordDF = pd.read_csv(recordFile)
    lastUpdt = recordDF['tradeDate'].values[-1]
    w.start()
    startDate = w.tdaysoffset(1, str(lastUpdt)).Data[0][0].strftime('%Y%m%d')
    if startDate>endDate:
        print('no new data to update')
        return
    betweenDays = w.tdays(startDate,endDate).Data[0]
    betweenDays = [int(tdt.strftime('%Y%m%d')) for tdt in betweenDays]
    print(betweenDays)
    outPut = {
        'tradeDate': [],
        'buyReturn': [],
        'holdReturn': [],
        'netReturn':[]
    }
    cols = ['stkcd', 'buyDate', 'ret', 'flag']
    for tdt in betweenDays:
        retList = pd.read_csv(os.path.join(holdListPath,'returns_tradeDate_{}.csv'.format(tdt)),encoding='gbk')
        retList = retList.loc[:, cols]
        validBuy = np.all([retList['buyDate']==tdt, np.isin(retList['flag'],(0,3,4))] ,axis=0)
        validHold = np.all([retList['buyDate']<tdt, np.isin(retList['flag'],(0,2,3,4))] ,axis=0)
        outPut['tradeDate'].append(tdt)
        outPut['buyReturn'].append(retList.loc[validBuy,'ret'].mean())
        outPut['holdReturn'].append(retList.loc[validHold,'ret'].mean())
        outPut['netReturn'].append(retList.loc[validBuy | validHold,'ret'].mean())
    outPut = pd.DataFrame(outPut).loc[:, ['tradeDate', 'buyReturn', 'holdReturn', 'netReturn']]
    indices = ['000001.SH','000300.SH','000905.SH','000016.SH']
    idxRets = pd.DataFrame(w.wsd(','.join(indices),'pct_chg',startDate,endDate).Data, columns=indices)
    idxRets = idxRets.loc[:]/100
    outPut = pd.concat([outPut, idxRets],axis=1)
    outPut.to_csv(recordFile,mode='a+',index=False,header=False)
    print('recorder updated')


def generate_doc(tdate=None, docPath=None,tradeListPath=None,recorderPath=None):
    tdate = dt.datetime.today().strftime('%Y%m%d') if tdate is None else tdate
    tdate = '20180809'
    docPath = r'.\reportDocs' if docPath is None else docPath
    recorderPath = r'.\\' if recorderPath is None else recorderPath
    tradeListPath = r'.\buyLists' if tradeListPath is None else tradeListPath
    # 创建空白文档
    document = Document()
    document.add_heading(r'策略模拟收益日报 {}'.format(tdate),0)
    document.add_paragraph()
    # 加入选出的股票
    # w.start()
    nextDate = 20180810  # w.tdaysoffset(1,tdate).Data[0][0].strftime('%Y%m%d')
    para = document.add_paragraph('')
    run = para.add_run(r'1. 选出股票列表：用于{}日开盘买入'.format(nextDate))
    run.font.bold = True
    selectStocks = pd.read_csv(os.path.join(tradeListPath,'stkNum50_tradeList_infoDate_{}.csv'.format(tdate)),encoding='gbk')
    selectStocks.sort_values(by=['stkcd'],inplace=True)
    selectStocks.index = range(selectStocks.shape[0])
    rows, cols = (25, 4)
    table = document.add_table(rows=rows, cols=cols, style='Table Grid')
    for rw in range(rows):
        table.cell(rw, 0).text = str(selectStocks['stkcd'][rw])
        table.cell(rw, 1).text = str(selectStocks['stkname'][rw])
        table.cell(rw, 2).text = str(selectStocks['stkcd'][rw + rows])
        table.cell(rw, 3).text = str(selectStocks['stkname'][rw + rows])

    # 计算 日度、周度、月度、年度 年化收益率，最大回撤，夏普比率
    document.add_paragraph()
    para = document.add_paragraph('')
    run = para.add_run(r'2. 模拟策略收益净值曲线: 费用设为双边千三')
    run.font.bold = True
    recorders = pd.read_csv(os.path.join(recorderPath, r'backtest_recorder.csv'))
    tdateWeekday = dt.datetime.strptime(tdate,'%Y%m%d').weekday()
    tradeDTs = recorders['tradeDate'].map(lambda x : dt.datetime.strptime(str(x),'%Y%m%d'))
    weekDays = tradeDTs.map(lambda x : x.weekday()).values
    idx = weekDays>=tdateWeekday
    idx[-1] = False
    lastWeekDate = recorders.loc[idx,'tradeDate'].values[-1]
    months = tradeDTs.map(lambda x : x.month)
    lastMonthDate = recorders.loc[months!=months.values[-1],'tradeDate'].values[-1]
    years = tradeDTs.map(lambda x : x.year)
    lastYearDate = recorders.loc[years != years.values[-1], 'tradeDate'].values[-1]
    cutDates = [lastWeekDate, lastMonthDate, lastYearDate]
    fees = 3/1000
    annCnt = 245
    digit = 2
    recorders['feeRet'] = recorders['netReturn'] - fees/2
    indicators = {
        'freq': ['当日','当周','当月','当年'],
        'cumRet': [recorders['feeRet'].values[-1]],
        'annRet': [''],
        'annVol': [''],
        'maxDD': [''],
        'sharpe': [''],
    }
    for tdt in cutDates:
        rets = recorders.loc[recorders['tradeDate']>tdt,'feeRet']
        indicators['cumRet'].append(rets.sum())
        indicators['annRet'].append(rets.mean() * annCnt)
        indicators['annVol'].append(rets.std() * np.sqrt(annCnt))
        indicators['sharpe'].append(indicators['annRet'][-1]/indicators['annVol'][-1])
        netVal = (1 + rets).cumprod()
        netMax = np.max([netVal.cummax().values,np.ones_like(netVal.values)],axis=0)    # 第一天算作本周、本月，其最大回测应该和1比
        indicators['maxDD'].append(np.min(netVal.values/netMax-1))
    table = document.add_table(rows=5, cols=6, style='Table Grid')
    table.cell(0, 0).text = ''
    table.cell(0, 1).text = '累计收益率'
    table.cell(0, 2).text = '年化收益率'
    table.cell(0, 3).text = '年化波动率'
    table.cell(0, 4).text = '最大回撤'
    table.cell(0, 5).text = '夏普比率'
    for rw in range(1, 5):
        table.cell(rw, 0).text = indicators['freq'][rw-1]
        table.cell(rw, 1).text = '{}%'.format(str(np.round(indicators['cumRet'][rw-1]*100, digit))) if indicators['cumRet'][rw-1] else indicators['cumRet'][rw-1]
        table.cell(rw, 2).text = '{}%'.format(str(np.round(indicators['annRet'][rw-1]*100, digit))) if indicators['annRet'][rw-1] else indicators['annRet'][rw-1]
        table.cell(rw, 3).text = '{}%'.format(str(np.round(indicators['annVol'][rw-1]*100, digit))) if indicators['annVol'][rw-1] else indicators['annVol'][rw-1]
        table.cell(rw, 4).text = '{}%'.format(str(np.round(indicators['maxDD'][rw-1]*100, digit))) if indicators['maxDD'][rw-1] else indicators['maxDD'][rw-1]
        table.cell(rw, 5).text = str(np.round(indicators['sharpe'][rw - 1],digit*2)) if indicators['sharpe'][rw-1] else indicators['sharpe'][rw-1]
    document.add_paragraph(r'注：年化收益计算默认每年245个交易日')

    # 画出 最近一年净值曲线
    document.add_paragraph()
    para = document.add_paragraph('')
    run = para.add_run(r'3. 最近一年模拟策略收益净值曲线:')
    run.font.bold = True
    rets = recorders.loc[recorders['tradeDate'] >= lastYearDate, ['tradeDate','feeRet','000001.SH','000300.SH','000905.SH']]
    firstIdx = rets.index.values[0]
    rets.loc[firstIdx, ['feeRet','000001.SH','000300.SH','000905.SH']] = 0
    netVals = (1 + rets.loc[:,['feeRet','000001.SH','000300.SH','000905.SH']]).cumprod()
    netVals.columns = ['模拟净值','上证指数','沪深300','中证500']
    mpl.rcParams['font.sans-serif'] = ['SimHei']
    fig = plt.figure(figsize=(20, 13))
    for col in netVals.columns:
        plt.plot(netVals[col],label=col,lw=2)
    xtickSteps = range(0,netVals.shape[0],5)
    plt.xticks(rotation=70,fontsize=20)
    plt.yticks(fontsize=20)
    plt.xticks(netVal.index.values[xtickSteps],rets['tradeDate'].values[xtickSteps])
    plt.legend(loc = 'upper left',fontsize=30)
    plt.title('最近一年净值曲线',fontsize=30)
    # plt.show()
    figPath = os.path.join(r'.\netvalFigures','netval_figure_{}.png'.format(tdate))
    plt.savefig(figPath)

    pic = document.add_picture(figPath, width=Inches(6))

    document.save(os.path.join(docPath,'策略模拟收益日报_{}.docx'.format(tdate)))




if __name__=='__main__':
    # get_all_from_list()

    # w.start()
    # all = w.wsd('000001.SH,000300.SH,000905.SH,000016.SH','pct_chg','20180601','20180808').Data
    # print(all)

    # update_recorder()
    # generate_return_report(20180809)
    generate_doc()